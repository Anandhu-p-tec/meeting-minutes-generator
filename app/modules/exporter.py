from inspect import signature

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.modules.schema import MeetingMinutes
import os

def export_meeting_minutes_to_docx(meeting_minutes: MeetingMinutes, output_file: str) -> None:
    """
    Export meeting minutes to a DOCX file using a formal layout, simulating an official template.

    Args:
        meeting_minutes (MeetingMinutes): The meeting data object.
        output_file (str): Path to the output DOCX file.
    """
    document = Document()

    # Set default font to Times New Roman, size 13
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)

    # Set page margins
    section = document.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # -------- HEADER SECTION --------
    # Create a 2-column table to simulate the left/right header
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    table.columns[0].width = Cm(8)
    table.columns[1].width = Cm(8)

    # Left cell: Organization name
    cell_left_1 = table.cell(0, 0).paragraphs[0]
    cell_left_1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_left_1 = cell_left_1.add_run("PARENT ORGANIZATION NAME\nORGANIZATION NAME\n-------")
    run_left_1.font.name = 'Times New Roman'
    run_left_1.font.size = Pt(13)
    run_left_1.bold = True

    # Right cell: Government slogan
    cell_right_1 = table.cell(0, 1).paragraphs[0]
    cell_right_1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_right_1 = cell_right_1.add_run("SOCIALIST REPUBLIC OF VIETNAM\nIndependence - Freedom - Happiness\n-------")
    run_right_1.font.name = 'Times New Roman'
    run_right_1.font.size = Pt(13)
    run_right_1.bold = True

    # Second row: Document number
    cell_left_2 = table.cell(1, 0).paragraphs[0]
    cell_left_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_left_2 = cell_left_2.add_run("")

    cell_right_2 = table.cell(1, 1).paragraphs[0]
    cell_right_2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_right_2 = cell_right_2.add_run("Document No: ....../MM-....")
    run_right_2.font.name = 'Times New Roman'
    run_right_2.font.size = Pt(13)

    # Spacer line
    document.add_paragraph("")

    # -------- MAIN TITLE --------
    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_para.add_run("MEETING MINUTES")
    run_title.font.name = "Times New Roman"
    run_title.font.size = Pt(16)
    run_title.bold = True

    document.add_paragraph("")

    # -------- MEETING DETAILS --------

    if meeting_minutes.meeting_time:
        p_time = document.add_paragraph()
        p_time.add_run("Start Time: ").bold = True
        p_time.add_run(meeting_minutes.meeting_time)

    if meeting_minutes.location:
        p_place = document.add_paragraph()
        p_place.add_run("Location: ").bold = True
        p_place.add_run(meeting_minutes.location)

    document.add_paragraph("")

    document.add_paragraph("Meeting Attendees: ", style="List Paragraph").runs[0].bold = True
    if meeting_minutes.host:
        document.add_paragraph(f"- Host: {meeting_minutes.host}", style="List Bullet")
    if meeting_minutes.note_taker:
        document.add_paragraph(f"- Note Taker: {meeting_minutes.note_taker}", style="List Bullet")
    if meeting_minutes.attendees:
        document.add_paragraph(f"- Participants: " + ", ".join(meeting_minutes.attendees), style="List Bullet")

    if meeting_minutes.meeting_goal:
        doc_para = document.add_paragraph()
        doc_para.add_run("Meeting Purpose: ").bold = True
        doc_para.add_run(meeting_minutes.meeting_goal)

    document.add_paragraph("")

    if meeting_minutes.agenda:
        doc_para = document.add_paragraph()
        doc_para.add_run("Agenda: ").bold = True
        for item in meeting_minutes.agenda:
            document.add_paragraph(item, 'List Bullet')
        document.add_paragraph("")

    if meeting_minutes.discussion_content:
        doc_para = document.add_paragraph()
        doc_para.add_run("Discussion Content: ").bold = True
        for topic, points in meeting_minutes.discussion_content.items():
            sub_heading = document.add_paragraph(topic)
            sub_heading.style = 'List Bullet'
            for point in points:
                document.add_paragraph(point, 'List Bullet 2')

    document.add_paragraph("")

    if meeting_minutes.decisions:
        doc_para = document.add_paragraph()
        doc_para.add_run("Decisions:").bold = True
        for decision in meeting_minutes.decisions:
            document.add_paragraph(decision, style='List Bullet')

    document.add_paragraph("")

    if meeting_minutes.conclusion:
        doc_para = document.add_paragraph()
        doc_para.add_run("Conclusion: ").bold = True
        doc_para.add_run(meeting_minutes.conclusion)

    document.add_paragraph("")

    if meeting_minutes.attached_documents:
        doc_para = document.add_paragraph()
        doc_para.add_run("Attachments: ").bold = True
        for item in meeting_minutes.attached_documents:
            document.add_paragraph(item, 'List Bullet')
    else:
        document.add_paragraph("Attachments: None")

    document.add_paragraph("")

    if meeting_minutes.notes:
        doc_para = document.add_paragraph()
        doc_para.add_run("Additional Notes: ").bold = True
        doc_para.add_run(meeting_minutes.notes)
    else:
        document.add_paragraph("Additional Notes: None")

    document.add_paragraph("")

    # -------- SIGNATURE SECTION --------
    signature_table = document.add_table(rows=1, cols=2)
    signature_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signature_table.columns[0].width = Cm(8)
    signature_table.columns[1].width = Cm(8)

    cell_left = signature_table.cell(0, 0).paragraphs[0]
    cell_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_left = cell_left.add_run("HOST\n\n\n\n\n(Sign and print full name)")
    run_left.font.name = "Times New Roman"
    run_left.font.size = Pt(13)

    cell_right = signature_table.cell(0, 1).paragraphs[0]
    cell_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_right = cell_right.add_run("NOTE TAKER\n\n\n\n\n(Sign and print full name)")
    run_right.font.name = "Times New Roman"
    run_right.font.size = Pt(13)

    document.save(output_file)


# --- Standalone Testing ---
if __name__ == "__main__":
    # Example meeting minutes data
    sample_minutes = MeetingMinutes(
        meeting_date="30/03/2025",
        meeting_time="10:00 - 11:30",
        location="Meeting Room A",
        host="John A. Smith",
        note_taker="Jane B. Tran",
        attendees=["John A. Smith", "Jane B. Tran", "Leo C. Le", "Dana D. Pham"],
        meeting_goal="Evaluate progress and plan next phase",
        agenda=["Progress report", "Challenges", "Solutions and assignments"],
        discussion_content={
            "Progress report": ["70% of work completed"],
            "Challenges": ["Not enough testers for QA phase"],
            "Solutions": ["Add 2 more testers", "Leo C. to lead", "Deadline: 15/04/2025"]
        },
        decisions=["Approve staff additions", "Reassign responsibilities"],
        conclusion="Finalize deadline and next meeting set for 20/04/2025",
        attached_documents=None,
        notes="Ensure task allocation and timeline tracking."
    )

    print("----- Sample Meeting Minutes -----")
    print(sample_minutes.model_dump())

    output_docx = "meeting_minutes_refined.docx"
    export_meeting_minutes_to_docx(sample_minutes, output_docx)
    print(f"Meeting minutes exported to file: {output_docx}")
